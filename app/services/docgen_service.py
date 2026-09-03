"""Генерация счёта и договора из принятого КП + отметки оплат (фаза 21).

Счёт — печатная форма HTML→PDF (print/invoice.html), договор — docx через
docxtpl. Оба фиксируются записями Document и видны в табе «Документы» лида.
"""

from datetime import datetime
from decimal import Decimal
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt
from num2words import num2words
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models import CompanyProfile, Deal, Document, Lead, Quote


def money_in_words(value: Decimal | float) -> str:
    """145200.00 → «Сто сорок пять тысяч двести рублей 00 копеек»."""
    d = Decimal(str(value)).quantize(Decimal("0.01"))
    rub = int(d)
    kop = int((d - rub) * 100)

    def plural(n: int, one: str, few: str, many: str) -> str:
        if n % 100 in (11, 12, 13, 14):
            return many
        if n % 10 == 1 and n % 100 != 11:
            return one
        if n % 10 in (2, 3, 4):
            return few
        return many

    words = num2words(rub, lang="ru")
    words = words[0].upper() + words[1:]
    return f"{words} {plural(rub, 'рубль', 'рубля', 'рублей')} {kop:02d} {plural(kop, 'копейка', 'копейки', 'копеек')}"


async def next_doc_number(session: AsyncSession, prefix: str) -> str:
    from app.services.quote_service import next_quote_number

    return await next_quote_number(session, prefix)


async def _ensure_deal(session: AsyncSession, quote: Quote) -> Deal:
    deal = (
        await session.execute(
            select(Deal).where(Deal.lead_id == quote.lead_id, Deal.status != "lost")
            .order_by(Deal.id.desc()).limit(1)
        )
    ).scalar_one_or_none()
    if deal is None:
        deal = Deal(lead_id=quote.lead_id, user_id=quote.user_id,
                    title=f"По КП {quote.number}", status="new")
        session.add(deal)
        await session.flush()
    return deal


async def get_or_create_invoice(session: AsyncSession, quote: Quote) -> Document:
    """Счёт по принятому КП. Повторный вызов возвращает существующий документ.
    Document не имеет quote_id, поэтому связь — по лиду+сумме+заголовку."""
    existing = (
        await session.execute(
            select(Document)
            .where(
                Document.doc_type == "invoice",
                Document.lead_id == quote.lead_id,
                Document.amount == float(quote.total),
                Document.title == f"Счёт на оплату по КП {quote.number}",
            )
            .order_by(Document.id.desc()).limit(1)
        )
    ).scalar_one_or_none()
    if existing:
        return existing

    deal = await _ensure_deal(session, quote)
    number = await next_doc_number(session, "СЧ")
    doc = Document(
        deal_id=deal.id,
        lead_id=quote.lead_id,
        user_id=quote.user_id,
        doc_type="invoice",
        title=f"Счёт на оплату по КП {quote.number}",
        number=number,
        amount=float(quote.total),
        status="draft",
    )
    session.add(doc)
    await session.flush()
    return doc


async def get_or_create_contract(session: AsyncSession, quote: Quote, profile: CompanyProfile) -> Document:
    """Договор по принятому КП: docx через docxtpl, файл в storage/documents/."""
    existing = (
        await session.execute(
            select(Document)
            .where(
                Document.doc_type == "contract",
                Document.lead_id == quote.lead_id,
                Document.title == f"Договор поставки по КП {quote.number}",
            )
            .order_by(Document.id.desc()).limit(1)
        )
    ).scalar_one_or_none()
    if existing:
        return existing

    deal = await _ensure_deal(session, quote)
    number = await next_doc_number(session, "Д")
    doc = Document(
        deal_id=deal.id,
        lead_id=quote.lead_id,
        user_id=quote.user_id,
        doc_type="contract",
        title=f"Договор поставки по КП {quote.number}",
        number=number,
        amount=float(quote.total),
        status="draft",
    )
    session.add(doc)
    await session.flush()

    lead = await session.get(Lead, quote.lead_id)
    doc_dir = Path("storage/documents") / str(quote.lead_id)
    doc_dir.mkdir(parents=True, exist_ok=True)
    doc_path = doc_dir / f"contract_{doc.id}.docx"
    render_contract_docx(quote, lead, profile, doc, doc_path)
    doc.file_path = str(doc_path)
    await session.flush()
    return doc


def ensure_contract_template(path: Path) -> Path:
    """Дефолтный docx-шаблон договора (docxtpl). Создаётся при первом
    использовании; владелец может заменить файл по этому пути своим."""
    if path.is_file():
        return path
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = DocxDocument()

    def p(text: str, bold: bool = False, size: int = 11) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    p("ДОГОВОР ПОСТАВКИ № {{ contract_number }} от {{ contract_date }}", bold=True)
    p("Поставщик: {{ supplier_name }}, ИНН {{ supplier_inn }}, {{ supplier_address }}, "
      "р/с {{ supplier_account }} в {{ supplier_bank }}, БИК {{ supplier_bic }}, "
      "в лице {{ supplier_director }}, действующего на основании Устава,")
    p("с одной стороны, и Покупатель: {{ buyer_name }}{% if buyer_inn %}, ИНН {{ buyer_inn }}{% endif %}"
      "{% if buyer_address %}, {{ buyer_address }}{% endif %}, с другой стороны, заключили настоящий договор:")
    p("1. Предмет договора", bold=True)
    p("Поставщик обязуется поставить, а Покупатель — принять и оплатить товар согласно "
      "спецификации (составляет неотъемлемую часть настоящего договора) на основании "
      "коммерческого предложения {{ quote_number }} от {{ quote_date }} на общую сумму "
      "{{ total_num }} ({{ total_words }}).")
    p("2. Цена и порядок оплаты", bold=True)
    p("Цена товара фиксируется в спецификации. Оплата производится на расчётный счёт "
      "Поставщика в порядке 100% предоплаты, если спецификацией не установлено иное. "
      "{{ payment_terms }}")
    p("3. Сроки поставки", bold=True)
    p("Срок поставки согласовывается сторонами в спецификации, но не превышает 30 "
      "календарных дней с момента поступления оплаты.")
    p("4. Прочие условия", bold=True)
    p("Во всём, что не урегулировано настоящим договором, стороны руководствуются "
      "действующим законодательством РФ. Договор вступает в силу с момента подписания.")
    p("5. Подписи сторон", bold=True)
    p("Поставщик: {{ supplier_director }} ___________________", )
    p("Покупатель: {{ buyer_signer or '____________________' }}")
    doc.save(path)
    return path


def render_contract_docx(quote: Quote, lead: Lead, profile: CompanyProfile,
                         doc: Document, out_path: Path) -> Path:
    """Заполняет docx-шаблон договора (docxtpl, Jinja-синтаксис в самом docx)."""
    from docxtpl import DocxTemplate

    tpl_path = ensure_contract_template(settings.STORAGE_DIR / "contract_template.docx")
    tpl = DocxTemplate(tpl_path)
    context = {
        "contract_number": doc.number,
        "contract_date": datetime.now().strftime("%d.%m.%Y"),
        "supplier_name": profile.name or "",
        "supplier_inn": profile.inn or "",
        "supplier_address": profile.legal_address or "",
        "supplier_account": profile.bank_account or "",
        "supplier_bank": profile.bank_name or "",
        "supplier_bic": profile.bank_bic or "",
        "supplier_director": profile.director_name or "",
        "buyer_name": lead.name,
        "buyer_inn": lead.inn or "",
        "buyer_address": lead.legal_address or lead.address or "",
        "buyer_signer": lead.head_name or "",
        "quote_number": quote.number,
        "quote_date": quote.created_at.strftime("%d.%m.%Y"),
        "total_num": f"{Decimal(str(quote.total)):,.2f}".replace(",", " ").replace(".", ","),
        "total_words": money_in_words(quote.total),
        "payment_terms": "",
    }
    tpl.render(context)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    tpl.save(out_path)
    return out_path


async def mark_invoice_paid(session: AsyncSession, doc: Document) -> None:
    """Оплата счёта → сделка paid + closed_at, лид стадия 6. Единственная точка
    синхронизации оплаты (кросс-фазное правило)."""
    if doc.doc_type != "invoice" or doc.status == "paid":
        return
    doc.status = "paid"
    doc.paid_at = datetime.now()
    doc.paid_amount = doc.amount
    if doc.deal_id:
        deal = await session.get(Deal, doc.deal_id)
        if deal:
            deal.status = "paid"
            deal.closed_at = datetime.now()
    lead = await session.get(Lead, doc.lead_id)
    if lead and lead.stage.isdigit() and int(lead.stage) < 6:
        from app.services.funnel_service import change_stage
        try:
            await change_stage(session, lead.id, "6", doc.user_id)
        except ValueError:
            pass
    await session.flush()

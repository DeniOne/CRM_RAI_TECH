import os
import re
from datetime import datetime

from docx import Document


DEFAULT_PLACEHOLDERS = {
    "company_name": lambda lead, user: lead.name,
    "inn": lambda lead, user: lead.inn or "",
    "district": lambda lead, user: lead.district or "",
    "settlement": lambda lead, user: lead.settlement or "",
    "address": lambda lead, user: lead.address or "",
    "head_name": lambda lead, user: lead.head_name or "",
    "site": lambda lead, user: lead.site or "",
    "region": lambda lead, user: lead.region.name if lead.region else "",
    "manager_name": lambda lead, user: user.full_name,
    "date": lambda lead, user: datetime.now().strftime("%d.%m.%Y"),
    "phone": lambda lead, user: lead.contacts[0].phone if lead.contacts else "",
    "email": lambda lead, user: lead.contacts[0].email if lead.contacts and lead.contacts[0].email else "",
    "rapeseed_volume": lambda lead, user: lead.rapeseed_volume or "",
    "harvest_timing": lambda lead, user: lead.harvest_timing or "",
}


def extract_placeholders(docx_path: str) -> list[str]:
    doc = Document(docx_path)
    placeholders = set()
    pattern = re.compile(r'\{(\w+)\}')

    for para in doc.paragraphs:
        placeholders.update(pattern.findall(para.text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    placeholders.update(pattern.findall(para.text))

    for section in doc.sections:
        for para in section.header.paragraphs:
            placeholders.update(pattern.findall(para.text))
        for para in section.footer.paragraphs:
            placeholders.update(pattern.findall(para.text))

    return sorted(placeholders)


def generate_document(template_path: str, replacements: dict, output_path: str) -> str:
    doc = Document(template_path)

    def replace_in_paragraphs(paragraphs):
        for para in paragraphs:
            full_text = para.text
            changed = False
            for key, val in replacements.items():
                placeholder = '{' + key + '}'
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, str(val))
                    changed = True
            if changed and para.runs:
                para.runs[0].text = full_text
                for run in para.runs[1:]:
                    run.text = ''

    replace_in_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_paragraphs(section.footer.paragraphs)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def convert_to_pdf(docx_path: str, pdf_path: str) -> str | None:
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return pdf_path
        return None
    except Exception:
        return None


def build_replacements(lead, user, extra: dict = None) -> dict:
    replacements = {}
    for key, func in DEFAULT_PLACEHOLDERS.items():
        try:
            replacements[key] = func(lead, user)
        except Exception:
            replacements[key] = ""
    if extra:
        for key, val in extra.items():
            if val:
                replacements[key] = val
    return replacements

from io import BytesIO

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from fastapi.responses import StreamingResponse


def render_xlsx(sheets: list[tuple[str, list[dict]]], filename: str = "report.xlsx") -> StreamingResponse:
    """
    sheets: список (sheet_name, rows). rows — list[dict] (или []).
    Каждая пара → отдельный лист xlsx.
    Возвращает StreamingResponse с корректными заголовками.
    """
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        for sheet_name, rows in sheets:
            df = pd.DataFrame(rows) if rows else pd.DataFrame()
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def render_docx(title: str, period: str, headers: list[str], rows: list[list], filename: str = "report.docx") -> StreamingResponse:
    """
    title: заголовок отчёта (add_heading, уровень 0).
    period: строка периода.
    headers: имена колонок (шапка таблицы, жирным).
    rows: list[list] — значения по строкам (без заголовка).
    """
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Период: {period}")

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Шапка — жирным
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Данные
    for row_idx, row in enumerate(rows):
        for col_idx, val in enumerate(row):
            table.rows[row_idx + 1].cells[col_idx].text = str(val) if val is not None else ""

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
